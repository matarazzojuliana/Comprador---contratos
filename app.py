# app.py
"""
Streamlit app para comparar PDF firmado vs Word original.
Genera un docx resaltando palabras agregadas/modificadas (rojo)
y palabras eliminadas (subrayadas en azul). Tambi�n muestra
un resumen de cambios e implicancias b�sicas.
"""
import streamlit as st
from pdf2docx import Converter
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import difflib, unidecode, re, tempfile, os, io
from collections import Counter

st.set_page_config(page_title="Comparador de Contratos", layout="centered")

# ---------- UI ----------
st.title("?? Comparador: contrato firmado vs. original")
st.markdown(
    "Sub� el **PDF firmado** y el **Word original**. La app convertir� el PDF a Word, "
    "comparar� los textos y generar� un Word con los cambios resaltados.\n\n"
    "**Rojo** = agregado/modificado � **Azul subrayado** = eliminado."
)

with st.expander("?? Notas importantes (leer)"):
    st.write("""
    - Si el PDF es un escaneo (imagen), la conversi�n no extraer� texto; se necesita OCR (Tesseract). 
    - Archivos muy grandes pueden tardar. 
    - Esto automatiza la revisi�n, pero cualquier decisi�n legal final debe validar el equipo legal.
    """)

pdf_file = st.file_uploader("1) Sub� el PDF firmado", type=["pdf"])
word_file = st.file_uploader("2) Sub� el Word original (.docx)", type=["docx"])

st.markdown("---")
col1, col2 = st.columns([1, 3])
with col1:
    use_ai = st.checkbox("Agregar an�lisis de implicancias mediante OpenAI (opcional)")
with col2:
    if use_ai:
        api_key = st.text_input("Peg� tu OpenAI API key (se usar� solo esta sesi�n)", type="password")
    else:
        api_key = None

compare_btn = st.button("?? Comparar documentos")

# ---------- helpers ----------
def add_underline(run):
    rPr = run._element.get_or_add_rPr()
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    u.set(qn('w:color'), '0000FF')
    rPr.append(u)

def normalize(text):
    text = unidecode.unidecode(text)
    text = text.lower()
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def extract_text_from_docx(path):
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return ' '.join(paragraphs)

def convert_pdf_to_docx(pdf_path, out_docx_path):
    cv = Converter(pdf_path)
    cv.convert(out_docx_path, start=0, end=None)
    cv.close()

def highlight_changes(original_doc_path, new_doc_path, output_path):
    """
    Basado en tu versi�n original (la que te funcion� mejor).
    Adem�s devuelve un resumen (counts, top palabras cambiadas).
    """
    text_old = extract_text_from_docx(original_doc_path)
    text_new = extract_text_from_docx(new_doc_path)

    norm_old = normalize(text_old)
    norm_new = normalize(text_new)

    sm = difflib.SequenceMatcher(None, norm_old.split(), norm_new.split())

    doc_out = Document()
    # leyenda
    legend = doc_out.add_paragraph()
    legend.add_run("?? Rojo: palabras agregadas o modificadas\n").font.color.rgb = RGBColor(255, 0, 0)
    run_blue = legend.add_run("?? Azul subrayado: palabras eliminadas (presentes en original, no en PDF firmado)\n")
    add_underline(run_blue)
    doc_out.add_paragraph()

    p = doc_out.add_paragraph()

    new_words_original = text_new.split()
    old_words_original = text_old.split()
    idx_new_word = 0
    idx_old_word = 0

    # resumen colecciones
    added_words = []
    deleted_words = []
    replaced_old = []
    replaced_new = []

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == 'equal':
            for _ in range(j2 - j1):
                if idx_new_word < len(new_words_original):
                    p.add_run(new_words_original[idx_new_word] + ' ')
                    idx_new_word += 1
                    idx_old_word += 1
        elif tag == 'replace':
            for _ in range(j2 - j1):
                if idx_new_word < len(new_words_original):
                    run = p.add_run(new_words_original[idx_new_word] + ' ')
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    replaced_new.append(new_words_original[idx_new_word])
                    idx_new_word += 1
            for _ in range(i2 - i1):
                if idx_old_word < len(old_words_original):
                    run = p.add_run(old_words_original[idx_old_word] + ' ')
                    add_underline(run)
                    replaced_old.append(old_words_original[idx_old_word])
                    idx_old_word += 1
        elif tag == 'insert':
            for _ in range(j2 - j1):
                if idx_new_word < len(new_words_original):
                    run = p.add_run(new_words_original[idx_new_word] + ' ')
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    added_words.append(new_words_original[idx_new_word])
                    idx_new_word += 1
        elif tag == 'delete':
            for _ in range(i2 - i1):
                if idx_old_word < len(old_words_original):
                    run = p.add_run(old_words_original[idx_old_word] + ' ')
                    add_underline(run)
                    deleted_words.append(old_words_original[idx_old_word])
                    idx_old_word += 1

    doc_out.save(output_path)

    # resumen estad�stico
    summary = {
        "counts": {
            "added": len(added_words),
            "deleted": len(deleted_words),
            "replaced_old": len(replaced_old),
            "replaced_new": len(replaced_new)
        },
        "added_top": Counter([normalize(w) for w in added_words]).most_common(15),
        "deleted_top": Counter([normalize(w) for w in deleted_words]).most_common(15),
        "replaced_old_top": Counter([normalize(w) for w in replaced_old]).most_common(15),
        "replaced_new_top": Counter([normalize(w) for w in replaced_new]).most_common(15),
        "raw_lists": {
            "added": added_words,
            "deleted": deleted_words,
            "replaced_old": replaced_old,
            "replaced_new": replaced_new
        }
    }
    return summary

def infer_implications_from_terms(terms):
    """
    Heur�stico: busca palabras clave en 'terms' y devuelve implicancias t�picas.
    No sustituye asesor�a legal.
    """
    keywords = {
        ("penaliz", "penalty", "multa"): "Aumenta exposici�n financiera por penalizaciones.",
        ("rescis", "rescisi�n", "resoluci"): "Reduce la capacidad de terminar el contrato anticipadamente.",
        ("plazo", "fecha", "termino", "vence"): "Modifica plazos; puede afectar entregables y SLA.",
        ("pago", "pagos", "factur"): "Impacta flujo de caja o condiciones de cobro/pago.",
        ("indemn", "indemniz"): "Aumenta potenciales obligaciones de indemnizaci�n.",
        ("confidenc", "confidential"): "Cambia condiciones de confidencialidad; riesgo de fuga de informaci�n.",
        ("jurisdic", "ley aplicable", "arbitra"): "Cambia la jurisdicci�n o mecanismo de resoluci�n de conflictos.",
        ("garant", "warranty", "garant�a"): "Modifica garant�as y responsabilidades por defectos."
    }
    implications = []
    norm_terms = " ".join([normalize(t) for t in terms])
    for keys, msg in keywords.items():
        for k in keys:
            if k in norm_terms:
                implications.append(msg)
                break
    return list(dict.fromkeys(implications))  # �nicos

# ---------- acci�n principal ----------
if compare_btn:
    if not pdf_file or not word_file:
        st.error("Sub� ambos archivos (PDF y Word) para comparar.")
    else:
        with st.spinner("Procesando... esto puede tardar unos segundos"):
            try:
                with tempfile.TemporaryDirectory() as tmp:
                    # Guardar archivos subidos
                    pdf_path = os.path.join(tmp, pdf_file.name)
                    with open(pdf_path, "wb") as f:
                        f.write(pdf_file.getbuffer())

                    word_path = os.path.join(tmp, word_file.name)
                    with open(word_path, "wb") as f:
                        f.write(word_file.getbuffer())

                    # Convertir PDF -> docx
                    converted_docx = os.path.join(tmp, "from_pdf.docx")
                    convert_pdf_to_docx(pdf_path, converted_docx)

                    # Comparar y generar doc final
                    output_docx = os.path.join(tmp, "comparado.docx")
                    summary = highlight_changes(word_path, converted_docx, output_docx)

                    # Lectura para descarga
                    with open(output_docx, "rb") as f:
                        data_bytes = f.read()

                    st.success("? Comparaci�n lista")
                    st.download_button(
                        label="?? Descargar DOCX con cambios",
                        data=data_bytes,
                        file_name=f"comparado_{os.path.splitext(pdf_file.name)[0]}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                    # Mostrar resumen
                    st.markdown("### ?? Resumen r�pido de cambios")
                    counts = summary["counts"]
                    st.write(f"- Palabras agregadas: **{counts['added']}**")
                    st.write(f"- Palabras eliminadas: **{counts['deleted']}**")
                    st.write(f"- Reemplazos (old/new): **{counts['replaced_old']} / {counts['replaced_new']}**")

                    st.markdown("#### ?? Top palabras agregadas")
                    st.table(summary["added_top"][:10] if summary["added_top"] else "�")

                    st.markdown("#### ?? Top palabras eliminadas")
                    st.table(summary["deleted_top"][:10] if summary["deleted_top"] else "�")

                    # Heur�stico de implicancias
                    raw_changed_terms = summary["raw_lists"]["added"] + summary["raw_lists"]["deleted"] + summary["raw_lists"]["replaced_new"] + summary["raw_lists"]["replaced_old"]
                    implications = infer_implications_from_terms(raw_changed_terms)
                    st.markdown("### ?? Implicancias (heur�sticas)")
                    if implications:
                        for imp in implications:
                            st.write(f"- {imp}")
                    else:
                        st.write("- No se detectaron t�rminos sensibles obvios. Revisar manualmente cl�usulas cr�ticas.")

                    # Opcional: an�lisis con OpenAI (si el usuario lo pidi� y peg� la API key)
                    if use_ai and api_key:
                        try:
                            st.markdown("### ?? An�lisis LLM (resumen + implicancias)")
                            # Llamada a OpenAI (requiere 'openai' en requirements y la API key)
                            import openai
                            openai.api_key = api_key

                            # Armamos prompt breve
                            prompt = (
                                "Act�a como analista legal de contratos. A partir de este breve resumen de cambios "
                                "entre el contrato original y el firmado, genera un p�rrafo corto explicando "
                                "las implicancias principales para la empresa contratante y dos recomendaciones de acci�n.\n\n"
                                f"Resumen estad�stico: {counts}\n"
                                f"Top agregadas: {summary['added_top'][:8]}\n"
                                f"Top eliminadas: {summary['deleted_top'][:8]}\n\n"
                                "Responde en espa�ol con bullets claros."
                            )

                            res = openai.ChatCompletion.create(
                                model="gpt-4o-mini",  # si no disponible, usar "gpt-4o" o "gpt-4"
                                messages=[{"role":"user","content":prompt}],
                                max_tokens=400,
                                temperature=0.2
                            )
                            text = res["choices"][0]["message"]["content"].strip()
                            st.write(text)
                        except Exception as e:
                            st.error(f"No se pudo ejecutar el an�lisis LLM: {e}")
                            st.info("Verific� que la API key y la disponibilidad del modelo sean correctas.")
            except Exception as e:
                st.error(f"Error durante el procesamiento: {e}")
                st.info("Si el PDF es escaneado, consider� usar OCR antes de comparar.")